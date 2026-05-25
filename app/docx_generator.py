import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.resume_schema import ApprovedResume

_BULLET = "•"
_NDASH = "–"


def _validated(raw: dict[str, Any]) -> dict[str, Any]:
    if "contact" in raw:
        return ApprovedResume.model_validate(raw).model_dump(by_alias=True)
    return ApprovedResume(
        contact={
            "name": raw.get("name", ""),
            "email": raw.get("email", ""),
            "phone": raw.get("phone", ""),
            "linkedin": raw.get("linkedin", ""),
            "github": raw.get("github", ""),
            "location": raw.get("location", ""),
        },
        summary=raw.get("summary", ""),
        experience=[{"title": "Experience", "bullets": raw.get("bullets", [])}],
        skills=raw.get("skills", []),
    ).model_dump(by_alias=True)


def generate_ats_docx(approved: dict[str, Any]) -> bytes:
    data = _validated(approved)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    contact = data["contact"]

    _name_para(doc, contact.get("name") or "")

    contact_parts = [v for v in [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        contact.get("linkedin"),
        contact.get("github"),
    ] if v]
    if contact_parts:
        _contact_line(doc, "  |  ".join(contact_parts))

    _hr(doc)

    if data.get("summary"):
        _section_heading(doc, "PROFESSIONAL SUMMARY")
        _body(doc, data["summary"])

    if data.get("experience"):
        _section_heading(doc, "EXPERIENCE")
        for i, item in enumerate(data["experience"]):
            _experience_block(doc, item)
            if i < len(data["experience"]) - 1:
                _gap(doc)

    if data.get("projects"):
        _section_heading(doc, "PROJECTS")
        for i, project in enumerate(data["projects"]):
            _project_block(doc, project)
            if i < len(data["projects"]) - 1:
                _gap(doc)

    if data.get("education"):
        _section_heading(doc, "EDUCATION")
        for edu in data["education"]:
            _education_block(doc, edu)

    if data.get("certifications"):
        _section_heading(doc, "CERTIFICATIONS")
        for cert in data["certifications"]:
            parts = [cert.get("name", ""), cert.get("issuer", ""), cert.get("year", "")]
            _body(doc, " — ".join(p for p in parts if p))

    if data.get("skills"):
        _section_heading(doc, "SKILLS")
        _body(doc, f"  {_BULLET}  ".join(data["skills"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _name_para(doc: Document, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(name)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)


def _contact_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _body(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def _bullet_para(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(f"{_BULLET}  {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def _meta_para(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _gap(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)


def _experience_block(doc: Document, item: dict[str, Any]) -> None:
    title = item.get("title", "")
    company = item.get("company", "")
    location = item.get("location", "")
    start_date = item.get("start_date", "")
    end_date = item.get("end_date", "")
    duration = item.get("duration", "")

    if start_date and end_date:
        date_str = f"{start_date} {_NDASH} {end_date}"
    elif start_date:
        date_str = f"{start_date} {_NDASH} Present"
    else:
        date_str = duration

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)

    title_run = p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    if company:
        sep = p.add_run("  |  ")
        sep.font.name = "Calibri"
        sep.font.size = Pt(10.5)
        sep.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        co = p.add_run(company)
        co.font.name = "Calibri"
        co.font.size = Pt(10.5)
        co.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    meta_parts = [v for v in [location, date_str] if v]
    if meta_parts:
        _meta_para(doc, "  |  ".join(meta_parts))

    for bullet_text in item.get("bullets", []):
        _bullet_para(doc, bullet_text)


def _project_block(doc: Document, project: dict[str, Any]) -> None:
    name = project.get("name", "")
    tech = ", ".join(project.get("technologies", []))
    description = project.get("description", "")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)

    name_run = p.add_run(name)
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(10.5)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    if tech:
        sep = p.add_run("  |  ")
        sep.font.name = "Calibri"
        sep.font.size = Pt(10)
        sep.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        tech_run = p.add_run(tech)
        tech_run.font.name = "Calibri"
        tech_run.font.size = Pt(10)
        tech_run.font.italic = True
        tech_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if description:
        _body(doc, description)

    for bullet_text in project.get("bullets", []):
        _bullet_para(doc, bullet_text)


def _education_block(doc: Document, edu: dict[str, Any]) -> None:
    degree = edu.get("degree", "")
    institution = edu.get("institution", "")
    location = edu.get("location", "")
    year = edu.get("year", "")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)

    deg_run = p.add_run(degree)
    deg_run.font.name = "Calibri"
    deg_run.font.size = Pt(10.5)
    deg_run.font.bold = True
    deg_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    if institution:
        sep = p.add_run("  |  ")
        sep.font.name = "Calibri"
        sep.font.size = Pt(10.5)
        sep.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        inst = p.add_run(institution)
        inst.font.name = "Calibri"
        inst.font.size = Pt(10.5)
        inst.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    meta_parts = [v for v in [location, year] if v]
    if meta_parts:
        _meta_para(doc, "  |  ".join(meta_parts))

    for detail in edu.get("details", []):
        _bullet_para(doc, detail)
