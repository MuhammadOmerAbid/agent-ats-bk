import re
from io import BytesIO
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from app.llm_client import extract_text_from_image


SKILL_TAXONOMY: dict[str, list[str]] = {
    "Python": ["python", "django", "fastapi", "flask"],
    "JavaScript": ["javascript", "js", "ecmascript"],
    "TypeScript": ["typescript", "ts"],
    "React": ["react", "react.js", "reactjs", "next.js", "nextjs"],
    "Node.js": ["node", "node.js", "express"],
    "REST APIs": ["rest api", "restful", "api development"],
    "GraphQL": ["graphql"],
    "PostgreSQL": ["postgresql", "postgres", "sql"],
    "MongoDB": ["mongodb", "mongo"],
    "Docker": ["docker", "containerization", "containers"],
    "Kubernetes": ["kubernetes", "k8s"],
    "AWS": ["aws", "amazon web services", "ec2", "s3", "lambda"],
    "Azure": ["azure"],
    "CI/CD": ["ci/cd", "continuous integration", "continuous deployment", "github actions"],
    "Git": ["git", "github", "gitlab"],
    "Testing": ["testing", "unit tests", "pytest", "jest", "cypress"],
    "Machine Learning": ["machine learning", "ml", "scikit-learn", "model training"],
    "NLP": ["nlp", "natural language processing", "spacy", "langchain"],
    "Data Analysis": ["data analysis", "pandas", "numpy", "analytics"],
    "Power BI": ["power bi", "dashboarding"],
    "Tailwind CSS": ["tailwind", "tailwind css"],
    "HTML/CSS": ["html", "css", "sass"],
}

SOFT_SKILLS = [
    "communication",
    "teamwork",
    "leadership",
    "problem solving",
    "collaboration",
    "ownership",
    "mentoring",
    "adaptability",
]

PREFERRED_HINTS = re.compile(r"\b(preferred|nice to have|bonus|plus|good to have|desirable)\b", re.I)
BULLET_PATTERN = re.compile(r"^\s*(?:[-*•]|\d+[.)])\s+(?P<text>.+)$")
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
JOB_SIGNAL_TERMS = [
    "job description",
    "responsibilities",
    "requirements",
    "qualifications",
    "we are hiring",
    "hiring",
    "apply",
    "position",
    "role",
    "candidate",
    "experience",
    "skills",
    "employment",
    "full-time",
    "part-time",
    "remote",
]


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()


def is_url(source: str) -> bool:
    return source.strip().lower().startswith(("http://", "https://"))


def extract_jd_text(source: str) -> str:
    normalized = source.strip()
    if not is_url(normalized):
        return normalized

    response = requests.get(
        normalized,
        timeout=10,
        headers={"User-Agent": "Mozilla/5.0 (compatible; AgentATS/0.1)"},
    )
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    return normalize_text(soup.get_text(separator=" ", strip=True))


def is_probably_job_description(text: str) -> bool:
    normalized = normalize_text(text).lower()
    word_count = len(normalized.split())
    if word_count < 50:
        return False

    signal_hits = sum(1 for term in JOB_SIGNAL_TERMS if term in normalized)
    skill_hits = len(extract_skills(normalized))
    has_job_title_pattern = bool(
        re.search(
            r"\b(engineer|developer|analyst|designer|manager|specialist|intern|consultant|architect|lead)\b",
            normalized,
        )
    )

    return signal_hits >= 3 or (signal_hits >= 2 and skill_hits >= 1) or (skill_hits >= 2 and has_job_title_pattern)


def extract_skills(text: str) -> list[str]:
    normalized = normalize_text(text)
    found = [
        canonical
        for canonical, aliases in SKILL_TAXONOMY.items()
        if any(_contains_alias(normalized, alias) for alias in aliases)
    ]
    return sorted(found)


def extract_jd_skills(jd_text: str) -> dict[str, list[str]]:
    normalized = normalize_text(jd_text)
    skills = extract_skills(normalized)
    preferred = [skill for skill in skills if _is_preferred_context(normalized, skill)]
    required = [skill for skill in skills if skill not in preferred]
    soft = [skill.title() for skill in SOFT_SKILLS if _contains_alias(normalized, skill)]

    return {"required": required, "preferred": preferred, "soft": soft}


def extract_resume_profile(resume_text: str) -> dict[str, object]:
    normalized = normalize_text(resume_text)
    bullets = extract_bullets(resume_text)

    return {
        "text": normalized,
        "skills": extract_skills(normalized),
        "bullets": bullets,
        "has_metrics": has_metrics(normalized),
        "word_count": len(normalized.split()),
    }


def extract_bullets(text: str) -> list[str]:
    bullets: list[str] = []
    for line in text.splitlines():
        match = BULLET_PATTERN.match(line)
        if match:
            bullets.append(match.group("text").strip())

    if bullets:
        return bullets[:24]

    sentences = re.split(r"(?<=[.!?])\s+", normalize_text(text))
    return [sentence.strip() for sentence in sentences if len(sentence.split()) >= 6][:24]


def has_metrics(text: str) -> bool:
    return bool(re.search(r"(\d+%|\$\d+|\b\d+\+?\s*(users|clients|projects|hours|seconds|days)\b)", text, re.I))


async def extract_text_from_upload(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = _clean_extracted_text(page.extract_text() or "")
            if page_text:
                pages.append(f"[Page {index}]\n{page_text}")
        return "\n\n".join(pages)

    if suffix in {".docx", ".doc"}:
        document = Document(BytesIO(content))
        return _extract_docx_text(document)

    if suffix in IMAGE_SUFFIXES:
        return extract_text_from_image(content, _mime_type_for_image(suffix))

    return content.decode("utf-8", errors="ignore")


def _extract_docx_text(document: Document) -> str:
    parts: list[str] = []

    for section in document.sections:
        _append_paragraphs(parts, section.header.paragraphs)
        _append_tables(parts, section.header.tables)

    _append_paragraphs(parts, document.paragraphs)
    _append_tables(parts, document.tables)

    for section in document.sections:
        _append_paragraphs(parts, section.footer.paragraphs)
        _append_tables(parts, section.footer.tables)

    return _clean_extracted_text("\n".join(parts))


def _append_paragraphs(parts: list[str], paragraphs: object) -> None:
    for paragraph in paragraphs:
        text = _clean_extracted_text(getattr(paragraph, "text", ""))
        if text:
            parts.append(text)


def _append_tables(parts: list[str], tables: object) -> None:
    for table in tables:
        for row in table.rows:
            cells = [_clean_extracted_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)
            for cell in row.cells:
                _append_tables(parts, cell.tables)


def _clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _mime_type_for_image(suffix: str) -> str:
    if suffix == ".png":
        return "image/png"
    if suffix == ".webp":
        return "image/webp"
    return "image/jpeg"


def _contains_alias(text: str, alias: str) -> bool:
    escaped = re.escape(alias.lower())
    return re.search(rf"(?<![a-z0-9]){escaped}(?![a-z0-9])", text.lower()) is not None


def _is_preferred_context(text: str, skill: str) -> bool:
    lower = text.lower()
    for alias in SKILL_TAXONOMY.get(skill, [skill]):
        index = lower.find(alias.lower())
        if index == -1:
            continue
        window = lower[max(0, index - 100) : index + len(alias) + 100]
        if PREFERRED_HINTS.search(window):
            return True
    return False
